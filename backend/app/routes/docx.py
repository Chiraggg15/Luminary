"""
DOCX Export Route  –  /api/docx
--------------------------------
POST /api/docx/generate  → Accept resume JSON and return a .docx file
"""

import io
from flask import Blueprint, request, jsonify, send_file
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from app.utils.jwt_helper import jwt_required_custom

docx_bp = Blueprint("docx", __name__)


@docx_bp.route("/generate", methods=["POST"])
@jwt_required_custom
def generate_docx():
    """Generate a professional DOCX resume from the provided resume data."""
    data = request.get_json() or {}

    personal = data.get("personal_info", {})
    experience = data.get("experience", [])
    education = data.get("education", [])
    skills = data.get("skills", {"technical": [], "soft": []})
    title = data.get("title", "Resume")

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    def add_heading_section(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x71, 0x71, 0x7A)  # zinc-500
        # Add a bottom border via XML
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'E4E4E7')
        pBdr.append(bottom)
        pPr.append(pBdr)
        return p

    # ── Name ───────────────────────────────────────────────────────────────
    name_para = doc.add_paragraph()
    name_para.paragraph_format.space_after = Pt(2)
    name_run = name_para.add_run((personal.get("full_name") or "YOUR NAME").upper())
    name_run.bold = True
    name_run.font.size = Pt(22)

    # ── Contact Line ───────────────────────────────────────────────────────
    contact_parts = [x for x in [
        personal.get("email"),
        personal.get("phone"),
        personal.get("location"),
        personal.get("linkedin"),
        personal.get("github"),
    ] if x]
    if contact_parts:
        contact_para = doc.add_paragraph(" • ".join(contact_parts))
        contact_para.paragraph_format.space_after = Pt(6)
        for run in contact_para.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x52, 0x52, 0x5B)

    # ── Summary ────────────────────────────────────────────────────────────
    if personal.get("summary"):
        add_heading_section("Summary")
        p = doc.add_paragraph(personal["summary"])
        p.paragraph_format.space_after = Pt(6)
        for run in p.runs:
            run.font.size = Pt(10)

    # ── Experience ─────────────────────────────────────────────────────────
    if experience:
        add_heading_section("Experience")
        for exp in experience:
            # Position + Company + Dates row
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(0)
            pos_run = p.add_run(exp.get("position", ""))
            pos_run.bold = True
            pos_run.font.size = Pt(11)
            p.add_run(f"  ·  {exp.get('company', '')}")
            p.runs[-1].font.size = Pt(10)
            p.runs[-1].font.color.rgb = RGBColor(0x0D, 0x9A, 0x6B)  # emerald

            # Dates
            date_str = f"{exp.get('start_date', '')} — {exp.get('end_date') or 'Present'}"
            date_para = doc.add_paragraph(date_str)
            date_para.paragraph_format.space_after = Pt(2)
            for run in date_para.runs:
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xA1, 0xA1, 0xAA)

            # Description bullets
            if exp.get("description"):
                for line in exp["description"].split("\n"):
                    line = line.strip()
                    if line:
                        bp = doc.add_paragraph(line, style="List Bullet")
                        bp.paragraph_format.space_after = Pt(1)
                        for run in bp.runs:
                            run.font.size = Pt(10)

    # ── Education ─────────────────────────────────────────────────────────
    if education:
        add_heading_section("Education")
        for edu in education:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            inst_run = p.add_run(edu.get("institution", ""))
            inst_run.bold = True
            inst_run.font.size = Pt(11)

            deg_parts = [e for e in [edu.get("degree"), edu.get("field_of_study")] if e]
            if deg_parts:
                dp = doc.add_paragraph(" in ".join(deg_parts))
                for run in dp.runs:
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0x52, 0x52, 0x5B)

            if edu.get("start_date") or edu.get("end_date"):
                dates = f"{edu.get('start_date', '')} — {edu.get('end_date', '')}"
                dp2 = doc.add_paragraph(dates)
                dp2.paragraph_format.space_after = Pt(2)
                for run in dp2.runs:
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0xA1, 0xA1, 0xAA)

    # ── Skills ────────────────────────────────────────────────────────────
    tech_skills = skills.get("technical", [])
    soft_skills = skills.get("soft", [])
    if tech_skills or soft_skills:
        add_heading_section("Skills")
        if tech_skills:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            label = p.add_run("Technical: ")
            label.bold = True
            label.font.size = Pt(10)
            val = p.add_run(", ".join(tech_skills))
            val.font.size = Pt(10)
        if soft_skills:
            p2 = doc.add_paragraph()
            p2.paragraph_format.space_after = Pt(2)
            label2 = p2.add_run("Soft Skills: ")
            label2.bold = True
            label2.font.size = Pt(10)
            val2 = p2.add_run(", ".join(soft_skills))
            val2.font.size = Pt(10)

    # ── Stream the file ────────────────────────────────────────────────────
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    filename = f"{title.replace(' ', '_')}.docx"

    return send_file(
        buffer,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
